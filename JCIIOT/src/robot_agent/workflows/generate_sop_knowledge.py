"""Generate contestant SOP Markdown from the official DOCX source files.

The workflow intentionally never reads ``knowledge/sop*.md``.  It combines:

* text and embedded images extracted from ``sop+prompt/*.docx``;
* VLM descriptions of the embedded images;
* read-only task/runtime metadata needed by the robot API; and
* a text LLM synthesis prompt retained in this source file.

Generated documents are written to ``team_submission/knowledge`` together
with a provenance manifest so judges can reproduce and inspect the process.
"""

from __future__ import annotations

import argparse
import base64
import hashlib
import json
import os
import re
import time
import urllib.request
from dataclasses import dataclass
from datetime import datetime, timezone
from pathlib import Path
from typing import Any

from docx import Document

from robot_agent.core.openai_client import OpenAIClient
_CASE_RE = re.compile(r"case\s+(\d+)", re.IGNORECASE)
_VISION_PROMPT = (
    "用一句话客观描述图中可见的工厂布局、工位、桌子、生产线、主要物体和清晰文字标签，"
    "不要推测看不见的信息，只输出最终答案。"
)
_VISION_RETRY_PROMPT = (
    "请直接输出一句中文最终答案：说明图中能看见的物体、工位或操作，并抄录清晰标签；"
    "不要解释思考过程。"
)

_SYNTHESIS_PROMPT = """You are generating an ORIGINAL contestant knowledge-base document
from an official Word SOP. Do not imitate or quote any prewritten reference Markdown.
Use only the supplied DOCX extraction, image observations, and verified runtime facts.

Write concise Markdown for a robot task planner. Requirements:
1. Preserve the material, visible Pick Station, visible Place Station, and quantity.
2. State the verified internal source, target, exact object names, and BC pose facts.
3. Give an executable sequence using only move, pick_up, and place_down. If
   runtime_object_port differs from semantic_source, move to runtime_object_port
   but keep pick_up.target equal to semantic_source for the competition event.
4. For quantity > 1, repeat the complete four-action cycle once per distinct object;
   never request an already transported object again.
5. State that pick_up is performed by the supplied BC policy and that transport
   attachment is valid only after BC grasp and lift verification both succeed.
6. Include navigation, grasp, placement, collision, drop, and anomaly constraints.
7. Clearly distinguish DOCX evidence from verified simulator/API metadata.
8. Do not claim facts absent from the supplied evidence.
9. Return Markdown only, with no code fence and no preamble.

The first deterministic execution-facts section will be added by the workflow,
so focus on a useful SOP explanation and validation checklist.

SOURCE DOCX: {source_name}
SOURCE SHA256: {source_sha256}

VERIFIED EXECUTION FACTS (read-only competition metadata):
{execution_facts}

DOCX TEXT:
{document_text}

VLM IMAGE OBSERVATIONS:
{image_observations}
"""


@dataclass(slots=True)
class ExtractedDocument:
    path: Path
    paragraphs: list[str]
    table_rows: list[list[str]]
    images: list[tuple[str, bytes]]

    @property
    def text(self) -> str:
        blocks = list(self.paragraphs)
        if self.table_rows:
            blocks.append("TABLE ROWS")
            blocks.extend(" | ".join(row) for row in self.table_rows)
        return "\n".join(blocks)


def _sha256_bytes(data: bytes) -> str:
    return hashlib.sha256(data).hexdigest()


def _atomic_write(path: Path, content: str) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    temporary = path.with_suffix(path.suffix + ".tmp")
    temporary.write_text(content, encoding="utf-8")
    temporary.replace(path)


def _extract_docx(path: Path) -> ExtractedDocument:
    document = Document(str(path))
    paragraphs = [p.text.strip() for p in document.paragraphs if p.text.strip()]
    table_rows: list[list[str]] = []
    for table in document.tables:
        for row in table.rows:
            values = [" ".join(cell.text.split()) for cell in row.cells]
            if any(values):
                table_rows.append(values)

    images_by_digest: dict[str, tuple[str, bytes]] = {}
    for relation in document.part.rels.values():
        if "image" not in relation.reltype:
            continue
        blob = relation.target_part.blob
        digest = _sha256_bytes(blob)
        name = relation.target_ref.split("/")[-1] if relation.target_ref else f"{digest[:12]}.png"
        images_by_digest.setdefault(digest, (name, blob))
    return ExtractedDocument(path, paragraphs, table_rows, list(images_by_digest.values()))


def _case_number(path: Path) -> int:
    match = _CASE_RE.search(path.name)
    if not match:
        raise ValueError(f"cannot determine case number from {path.name}")
    return int(match.group(1))


def _task_for_case(task_config: dict, case_number: int) -> dict:
    task_index = (case_number - 1) // 2
    tasks = task_config.get("tasks", [])
    if task_index < 0 or task_index >= len(tasks):
        raise ValueError(f"case {case_number} has no task_config entry")
    return dict(tasks[task_index])


def _runtime_inventory(task: dict) -> list[dict[str, Any]]:
    """Read object inventory from a headless scene and close it immediately."""
    from robot_agent.environments import RobosuiteBackend

    backend = RobosuiteBackend(
        env_name=str(task["env_name"]),
        camera="birdview",
        drive_mode="direct",
        headless=True,
    )
    try:
        backend.reset()
        raw = backend.env
        metadata = getattr(raw, "material_metadata", {}) or {}
        inventory: list[dict[str, Any]] = []
        for object_name, info in sorted(metadata.items()):
            if not isinstance(info, dict):
                continue
            fixed_pose = info.get("fixed_pose")
            if fixed_pose is None:
                fixed_pose = []
            inventory.append({
                "object_name": str(object_name),
                "port_name": str(info.get("port_name") or ""),
                "kind": str(info.get("kind") or ""),
                "fixed_xyz": [round(float(value), 4) for value in fixed_pose[:3]],
            })
        return inventory
    finally:
        backend.close()


def _execution_facts(task: dict, task_config: dict, inventory: list[dict[str, Any]]) -> dict:
    expected = str(task.get("object") or "")
    expected_info = next((item for item in inventory if item["object_name"] == expected), None)
    relevant_port = (expected_info or {}).get("port_name") or str(task["source"])
    expected_kind = (expected_info or {}).get("kind") or ""
    peers = [
        item["object_name"]
        for item in inventory
        if item["object_name"] == expected
        or (
            expected_kind
            and item.get("kind") == expected_kind
            and item.get("port_name") == relevant_port
        )
    ]
    if expected and expected not in peers:
        peers.insert(0, expected)
    # Single-object tasks must not present co-located, differently coloured
    # objects as alternatives to the configured target.  L5 intentionally
    # needs all three distinct same-kind objects.
    if str(task.get("level")) != "L5":
        peers = [expected] if expected else []

    grasp_pose = task_config.get("grasp_poses", {}).get(str(task["source"]), {})
    runtime_pose = task_config.get("grasp_poses", {}).get(str(relevant_port), {})
    return {
        "level": task.get("level"),
        "environment": task.get("env_name"),
        "semantic_source": task.get("source"),
        "semantic_target": task.get("target"),
        "configured_target_object": expected,
        "runtime_object_port": relevant_port,
        "source_bc_pose": grasp_pose,
        "runtime_object_bc_pose": runtime_pose,
        "same_kind_objects_at_runtime_port": peers,
    }


def _describe_images(
    extracted: ExtractedDocument,
    *,
    base_url: str,
    api_key: str,
    model: str,
) -> list[dict[str, str]]:
    descriptions: list[dict[str, str]] = []
    for index, (name, image) in enumerate(extracted.images, start=1):
        description = ""
        prompts = (_VISION_PROMPT, _VISION_RETRY_PROMPT, _VISION_RETRY_PROMPT)
        for attempt, prompt in enumerate(prompts, start=1):
            description = _ask_visual_caption(
                prompt,
                image,
                base_url=base_url,
                api_key=api_key,
                model=model,
            )
            if description.strip():
                break
            if attempt < len(prompts):
                time.sleep(float(attempt))
        if not description.strip():
            raise RuntimeError(
                f"VLM returned empty content three times for image {index} "
                f"in {extracted.path.name}"
            )
        descriptions.append({
            "name": name,
            "sha256": _sha256_bytes(image),
            "description": str(description).strip(),
        })
    return descriptions


def _ask_visual_caption(
    prompt: str,
    image: bytes,
    *,
    base_url: str,
    api_key: str,
    model: str,
) -> str:
    """Call the VLM with thinking disabled for a short factual caption.

    GLM-5V-Turbo otherwise may spend the complete output allowance in
    ``reasoning_content`` and return an empty final ``content`` field.  The
    caption is a direct perception task, so disabling reasoning is both faster
    and less ambiguous than treating hidden reasoning as the final answer.
    """
    if image.startswith(b"\x89PNG"):
        mime = "image/png"
    elif image.startswith(b"\xff\xd8"):
        mime = "image/jpeg"
    elif image.startswith(b"RIFF") and image[8:12] == b"WEBP":
        mime = "image/webp"
    else:
        mime = "image/png"

    payload = {
        "model": model,
        "messages": [{
            "role": "user",
            "content": [
                {"type": "text", "text": prompt},
                {
                    "type": "image_url",
                    "image_url": {
                        "url": (
                            f"data:{mime};base64,"
                            f"{base64.b64encode(image).decode('ascii')}"
                        ),
                    },
                },
            ],
        }],
        "thinking": {"type": "disabled"},
        "max_tokens": 1024,
        "stream": False,
    }
    request = urllib.request.Request(
        f"{base_url.rstrip('/')}/chat/completions",
        data=json.dumps(payload).encode("utf-8"),
        headers={
            "Authorization": f"Bearer {api_key}",
            "Content-Type": "application/json",
        },
    )
    response = json.loads(
        urllib.request.urlopen(request, timeout=180.0).read().decode("utf-8"),
    )
    return str(response["choices"][0]["message"].get("content") or "").strip()


def _clean_markdown(text: str) -> str:
    cleaned = text.strip()
    if cleaned.startswith("```"):
        cleaned = re.sub(r"^```(?:markdown)?\s*", "", cleaned, flags=re.IGNORECASE)
        cleaned = re.sub(r"\s*```$", "", cleaned)
    return cleaned.strip()


def _deterministic_header(
    *,
    case_number: int,
    source_name: str,
    source_sha256: str,
    task: dict,
    facts: dict,
    text_model: str,
    vision_model: str,
) -> str:
    peers = facts["same_kind_objects_at_runtime_port"]
    quantity_match = re.search(
        r"(?:quantity\s*(?:to transport)?\s*[:=]\s*|move\s+the\s+)(\d+|three|two|one)",
        task.get("source_prompt", ""),
        re.IGNORECASE,
    )
    quantity = quantity_match.group(1) if quantity_match else (3 if case_number == 9 else 1)
    runtime_source = facts["runtime_object_port"]
    semantic_source = task["source"]
    if str(task.get("level")) == "L5":
        object_rule = (
            "one cycle per distinct object: "
            + ", ".join(f"`{name}`" for name in peers)
        )
        cycle = (
            f"`move({runtime_source}) → pick_up(target={semantic_source}, "
            f"object_name=<next distinct object>) → move({task['target']}) "
            f"→ place_down({task['target']})`"
        )
    else:
        object_rule = f"exact object: `{task['object']}`"
        cycle = (
            f"`move({runtime_source}) → pick_up(target={semantic_source}, "
            f"object_name={task['object']}) → move({task['target']}) "
            f"→ place_down({task['target']})`"
        )
    return f"""# AI-Generated SOP — Case {case_number}

## Planner-Critical Facts

- `{task['level']}`; quantity **{quantity}**; {object_rule}
- Required BC cycle: {cycle}
- `pick_up` uses the supplied BC policy; transport is valid only after BC grasp and lift both succeed.

> Generated automatically from `{source_name}`; this is contestant output, not a copied reference SOP.
> Source SHA256: `{source_sha256}` · Text model: `{text_model}` · Vision model: `{vision_model}`

## Execution Facts

- Level: `{task['level']}`; environment: `{task['env_name']}`
- Task quantity: **{quantity}**
- Internal source: `{task['source']}`; internal target: `{task['target']}`
- Configured target object: `{task['object']}`
- Runtime object port: `{facts['runtime_object_port']}`
- Distinct eligible objects at that port: {', '.join(f'`{name}`' for name in peers) or 'none reported'}
- Grasp implementation: supplied **BC policy**, followed by lift verification
- Transport attachment rule: enable only after both BC grasp and lift verification succeed
"""


def generate_one(
    path: Path,
    *,
    task_config: dict,
    output_dir: Path,
    text_client: OpenAIClient,
    vlm_base_url: str,
    vlm_api_key: str,
    vlm_model: str,
) -> dict[str, Any]:
    started = time.perf_counter()
    case_number = _case_number(path)
    task = _task_for_case(task_config, case_number)
    extracted = _extract_docx(path)
    if extracted.paragraphs:
        task["source_prompt"] = extracted.paragraphs[0]
    inventory = _runtime_inventory(task)
    facts = _execution_facts(task, task_config, inventory)
    images = _describe_images(
        extracted,
        base_url=vlm_base_url,
        api_key=vlm_api_key,
        model=vlm_model,
    )

    source_bytes = path.read_bytes()
    source_sha256 = _sha256_bytes(source_bytes)
    prompt = _SYNTHESIS_PROMPT.format(
        source_name=path.name,
        source_sha256=source_sha256,
        execution_facts=json.dumps(facts, ensure_ascii=False, indent=2),
        document_text=extracted.text,
        image_observations=json.dumps(images, ensure_ascii=False, indent=2),
    )
    body = _clean_markdown(
        text_client.generate(prompt, num_predict=5000, temperature=0.1),
    )
    if not body:
        raise RuntimeError(f"text LLM returned empty Markdown for {path.name}")

    header = _deterministic_header(
        case_number=case_number,
        source_name=path.name,
        source_sha256=source_sha256,
        task=task,
        facts=facts,
        text_model=text_client.model,
        vision_model=vlm_model,
    )
    output_path = output_dir / f"generated_sop_case_{case_number}.md"
    _atomic_write(output_path, header.rstrip() + "\n\n" + body.rstrip() + "\n")
    return {
        "case": case_number,
        "source": str(path),
        "source_sha256": source_sha256,
        "output": str(output_path),
        "paragraphs": len(extracted.paragraphs),
        "table_rows": len(extracted.table_rows),
        "images": len(extracted.images),
        "images_described": sum(
            item["description"] != "VLM returned no description" for item in images
        ),
        "execution_facts": facts,
        "elapsed_sec": round(time.perf_counter() - started, 3),
    }


def main() -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--source-dir", type=Path, default=Path("sop+prompt"))
    parser.add_argument(
        "--output-dir", type=Path, default=Path("team_submission/knowledge"),
    )
    parser.add_argument("--task-config", type=Path, default=Path("knowledge/task_config.json"))
    parser.add_argument("--text-model", default=os.getenv("OPENAI_MODEL", "glm-5.2"))
    parser.add_argument("--vision-model", default=os.getenv("VLM_MODEL", "glm-5v-turbo"))
    parser.add_argument(
        "--cases",
        type=int,
        nargs="+",
        help="Generate only these official case numbers (for example: --cases 3 5).",
    )
    args = parser.parse_args()

    text_api_key = os.getenv("OPENAI_API_KEY", "")
    vlm_api_key = os.getenv("VLM_API_KEY", "") or text_api_key
    text_base_url = os.getenv("OPENAI_BASE_URL", "")
    vlm_base_url = os.getenv("VLM_BASE_URL", "") or text_base_url
    if not text_api_key or not text_base_url:
        raise SystemExit("OPENAI_API_KEY and OPENAI_BASE_URL are required")
    if not vlm_api_key or not vlm_base_url:
        raise SystemExit("VLM_API_KEY/VLM_BASE_URL (or OPENAI fallbacks) are required")

    task_config = json.loads(args.task_config.read_text(encoding="utf-8"))
    sources = sorted(args.source_dir.glob("JCIIOT 2026 case *.docx"))
    if args.cases:
        requested_cases = set(args.cases)
        sources = [source for source in sources if _case_number(source) in requested_cases]
        found_cases = {_case_number(source) for source in sources}
        missing_cases = sorted(requested_cases - found_cases)
        if missing_cases:
            raise SystemExit(f"official DOCX files not found for cases: {missing_cases}")
    elif len(sources) != 5:
        raise SystemExit(f"expected 5 official DOCX files, found {len(sources)}")

    text_client = OpenAIClient(
        api_key=text_api_key,
        base_url=text_base_url,
        model=args.text_model,
        timeout=240.0,
    )
    generated = []
    for source in sources:
        print(f"GENERATE {source.name}", flush=True)
        item = generate_one(
            source,
            task_config=task_config,
            output_dir=args.output_dir,
            text_client=text_client,
            vlm_base_url=vlm_base_url,
            vlm_api_key=vlm_api_key,
            vlm_model=args.vision_model,
        )
        generated.append(item)
        print(
            f"DONE case={item['case']} paragraphs={item['paragraphs']} "
            f"images={item['images_described']}/{item['images']} "
            f"elapsed={item['elapsed_sec']}s output={item['output']}",
            flush=True,
        )

    manifest_path = args.output_dir / "generated_sop_manifest.json"
    prior_documents: list[dict[str, Any]] = []
    if args.cases and manifest_path.exists():
        prior_manifest = json.loads(manifest_path.read_text(encoding="utf-8"))
        prior_documents = [
            item
            for item in prior_manifest.get("documents", [])
            if int(item.get("case", -1)) not in set(args.cases)
        ]

    manifest = {
        "generated_at_utc": datetime.now(timezone.utc).isoformat(),
        "generator": "robot_agent.workflows.generate_sop_knowledge",
        "text_model": args.text_model,
        "vision_model": args.vision_model,
        "reference_markdown_read": False,
        "documents": sorted(prior_documents + generated, key=lambda item: int(item["case"])),
    }
    _atomic_write(manifest_path, json.dumps(manifest, ensure_ascii=False, indent=2) + "\n")
    print(f"MANIFEST {manifest_path}", flush=True)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
